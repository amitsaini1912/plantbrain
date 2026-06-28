"""
make_samples.py
---------------
Generates a set of realistic industrial documents for demo purposes.
Run once: python make_samples.py
"""

from pathlib import Path
import fitz       # PyMuPDF
from docx import Document

OUT = Path("sample_docs")
OUT.mkdir(exist_ok=True)


# ═══════════════════════════════════════════════════════════════════════════════
# 1. Pump P-101 Maintenance & Operations Manual (PDF, 4 pages)
# ═══════════════════════════════════════════════════════════════════════════════
PUMP_PAGES = [
    """\
PUMP P-101 – CENTRIFUGAL FEED PUMP
MAINTENANCE & OPERATIONS MANUAL  Rev 4  |  Jan 2025

Equipment Tag : P-101 (Spare: P-101S)
Location      : Unit 2 – Coker Feed Section, Platform Level EL+12.0 m
Manufacturer  : KSB India Ltd  (Model: Etanorm SYT 150-315)
Rated Flow    : 450 m³/h   Rated Head : 120 m   NPSH required: 4.2 m
Motor         : 250 kW, 2-pole, FLP enclosure, IE3 efficiency class
Fluid         : Coker feed (Vacuum residue), SG 0.92, 185 °C at inlet

1. OVERVIEW
P-101 delivers heated feedstock from the surge drum D-204 to fired heater H-301.
It is a critical-service pump operating 24×7. Standby P-101S must remain on
auto-standby and should be alternated every 72 hours per the rotating equipment
policy. Any trip of P-101 automatically starts P-101S within 8 seconds.
""",
    """\
2. ROUTINE MAINTENANCE SCHEDULE

Daily checks (operator, each shift):
  • Seal flush flow indicator: minimum 3 L/min; low flow triggers alarm PA-2101
  • Bearing temperature: alarm 85 °C, trip 95 °C (bearing RTDs TI-2101A/B)
  • Vibration: alarm 7.1 mm/s, trip 11.2 mm/s (transmitters VT-2101A/B)
  • Suction pressure: minimum 2.5 barg at PI-2101; below 2.1 barg causes cavitation

Weekly (area technician):
  • Inspect mechanical seal for weeping; any visible drip > 5 drops/min → WO request
  • Record suction (PI-2101) and discharge (PI-2102) pressures in log sheet FM-M-014

Monthly:
  • Grease motor bearings: Kluber Isoflex NBU15 grease, 3 shots per side; do NOT over-grease
  • Check coupling alignment; maximum allowable: 0.10 mm parallel, 0.05 mm angular

Annual Overhaul:
  • Pull pump for full teardown; replace mechanical seal (Item 433), wear rings, and bearings
  • Laser-align coupling within 0.05 mm
  • Hydraulic test casing at 1.5× operating pressure

Known failure mode – Cavitation-induced seal failure:
Repeated seal failures on P-101 (WO#8821, WO#9034, WO#9412) have been traced to
cavitation when D-204 level falls below 40%. Always confirm drum level > 40% before
ramping throughput above 380 m³/h. Install low-level interlock LI-2204 per RCA#2024-07.
""",
    """\
3. SAFETY PRECAUTIONS

P-101 handles hydrocarbon at 185 °C, well above auto-ignition temperature (AIT ~250 °C).

Pre-Maintenance Mandatory Steps:
  1. Obtain Maintenance Work Permit (MWP) signed by Area Authority and Safety Officer
  2. Lock out and tag out (LOTO) motor at MCC panel – minimum 1 padlock per worker
  3. Isolate suction (XV-2101) and discharge (XV-2102) block valves; install spades
  4. Depressurise to closed blowdown header (BD-03); verify zero pressure at PI-2101
  5. Drain and flush casing; confirm skin temp < 40 °C before touching
  6. For pump pit access: gas test LEL < 10%, O₂ 19.5–23.5%, H₂S < 10 ppm
  7. Issue Confined Space Entry Permit (CSEP) – do NOT enter without valid CSEP

Hot-Work Restriction:
No hot-work permit (HWP) may be active within 15 metres of P-101 during any
maintenance involving casing opening. SIMOPS (Simultaneous Operations) Procedure
OISD-GDN-192 applies. Inform shift supervisor before issuing HWP in adjacent area.

Do NOT dead-head the pump (run against closed discharge valve) for more than 30 seconds —
thermal runaway occurs within 90 seconds and can rupture the seal.
""",
    """\
4. TROUBLESHOOTING GUIDE

Symptom: High vibration (> 7.1 mm/s)
  Possible causes:
    a. Cavitation — check suction pressure; confirm D-204 level > 40%
    b. Bearing wear — pull bearing for inspection; replace if pitting or spalling
    c. Coupling misalignment — laser-align; tolerance ≤ 0.10 mm

Symptom: Seal leakage
  Possible causes:
    a. Cavitation damage to seal faces — inspect and replace seal
    b. Flush flow inadequate — clean strainer in flush line; min 3 L/min
    c. Shaft run-out > 0.05 mm — check shaft straightness; replace if bent

Symptom: Low discharge pressure
  Possible causes:
    a. Wear rings eroded — measure clearance; replace if > 1.5× OEM spec
    b. Suction strainer fouled — schedule clean-out on next shutdown
    c. Wrong rotation direction — verify after any motor rewind

5. SPARE PARTS LIST (CRITICAL SPARES – MAINTAIN ON-SITE)
  Item 433  Mechanical seal (EagleBurgmann BRT-D) – 2 sets
  Item 321  Radial bearing (SKF 6215) – 4 units
  Item 325  Thrust bearing (SKF 7315) – 2 units
  Item 201  Wear ring set – 1 set
  Item 501  Coupling element (Rexnord M90) – 2 units
""",
]

doc = fitz.open()
for body in PUMP_PAGES:
    page = doc.new_page(width=595, height=842)
    page.insert_textbox(
        fitz.Rect(50, 50, 545, 800), body,
        fontsize=10, fontname="helv", align=0
    )
doc.save(OUT / "pump_P-101_manual.pdf")
doc.close()
print("[OK] pump_P-101_manual.pdf")


# ═══════════════════════════════════════════════════════════════════════════════
# 2. Confined Space Entry SOP (Word .docx)
# ═══════════════════════════════════════════════════════════════════════════════
d = Document()
d.add_heading("Confined Space Entry – Standard Operating Procedure", level=1)
d.add_paragraph(
    "Document No: SOP-SAF-014   Revision: 4   Effective: 01-Mar-2025\n"
    "Owner: Plant Safety Department   Approver: Head of HSE"
)

d.add_heading("1. Purpose", level=2)
d.add_paragraph(
    "This procedure defines mandatory controls for entry into any confined space — "
    "including pump pits, vessels (V-series), tanks (TK-series), columns (C-series), "
    "and below-grade pipelines — to prevent fatalities from toxic or oxygen-deficient atmospheres. "
    "It is mandatory under the Factories Act 1948 (Sec 36, 36A) and OISD-STD-116."
)

d.add_heading("2. Scope", level=2)
d.add_paragraph(
    "Applies to all contractors and employees entering any space classified as confined space "
    "in the Site Confined Space Register (HSE-REG-CS-2024)."
)

d.add_heading("3. Atmospheric Testing Requirements", level=2)
d.add_paragraph(
    "The following limits must be confirmed by a certified gas tester before every entry:\n"
    "  • Oxygen:          19.5 % – 23.5 %\n"
    "  • Flammable Gas:   < 10 % LEL  (alarm at 5 % LEL)\n"
    "  • Hydrogen Sulphide (H₂S): < 10 ppm TWA, < 15 ppm ceiling\n"
    "  • Carbon Monoxide (CO): < 35 ppm TWA\n"
    "Testing shall be continuous during entry using calibrated 4-gas detector."
)

d.add_heading("4. Mandatory Controls", level=2)
d.add_paragraph(
    "Before issuing a Confined Space Entry Permit (CSEP):\n"
    "  1. Area Authority signs the CSEP; copies to Safety Dept and Shift Supervisor\n"
    "  2. Equipment isolated, de-energised, and LOTO applied (see LOTO-SOP-015)\n"
    "  3. Vessel purged and ventilated; forced air supply during entry\n"
    "  4. Trained standby attendant posted outside — may not enter under any circumstance\n"
    "  5. Rescue equipment (tripod, harness, SCBA) positioned at entry point\n"
    "  6. Entrant wears personal 4-gas detector; evacuation if any alarm activates"
)

d.add_heading("5. SIMOPS Restrictions", level=2)
d.add_paragraph(
    "Simultaneous Operations (SIMOPS) Policy per OISD-GDN-192:\n"
    "  • No hot-work permit (HWP) within 15 metres of any open confined space\n"
    "  • No pressure testing on adjacent lines while confined space is occupied\n"
    "  • Inform SIMOPS coordinator 1 hour before entry; log in Daily SIMOPS Register\n\n"
    "The 2024 site incident review (INC-2024-03) found that combined maintenance on "
    "P-101 and H-301 without a SIMOPS check preceded the most serious near-miss of the year. "
    "The investigation recommended mandatory SIMOPS briefing before every CSEP."
)

d.add_heading("6. Rescue Plan", level=2)
d.add_paragraph(
    "Non-entry rescue is the default (mechanical retrieval via tripod/harness). "
    "Entry rescue may only be initiated by a trained rescuer with SCBA; "
    "alert plant emergency response team (ERT) before entry rescue commences. "
    "Call emergency: Ext 100 (ERT), Ext 101 (Ambulance), Ext 102 (Fire Station)."
)
d.save(OUT / "SOP-SAF-014_confined_space.docx")
print("[OK] SOP-SAF-014_confined_space.docx")


# ═══════════════════════════════════════════════════════════════════════════════
# 3. Maintenance Work Order Log (plain text, simulates a CMMS export)
# ═══════════════════════════════════════════════════════════════════════════════
WO_LOG = """\
MAINTENANCE WORK ORDER LOG – UNIT 2 COKER FEED SECTION
Export: CMMS v4.2  |  Period: Jan 2023 – Dec 2024

WO#     | Date       | Equipment | Priority | Type      | Description                                   | Status   | Technician
--------|------------|-----------|----------|-----------|-----------------------------------------------|----------|------------
WO#8821 | 2023-02-14 | P-101     | HIGH     | CM        | Mechanical seal failure; detected by seal pot  | Closed   | Rajan K
                                                           overflow. Root cause: cavitation. D-204 level
                                                           was 32% at time of failure. Replaced seal
                                                           (Item 433). Recommended install LI-2204 low-
                                                           level interlock.
WO#8890 | 2023-03-21 | D-204     | MED      | PM        | Annual vessel inspection per IBR. Thickness    | Closed   | TPIA – Bureau Veritas
                                                           readings within tolerance. Corrosion allowance
                                                           remaining: 3.2 mm. Next inspection due 2025-03.
WO#9034 | 2023-07-09 | P-101     | HIGH     | CM        | Second seal failure within 5 months. Vibration | Closed   | Rajan K / Singh A
                                                           VT-2101A reading 8.9 mm/s prior to failure.
                                                           Bearing also replaced (Item 321). RCA initiated
                                                           (see RCA#2023-04). Suspected cavitation again.
WO#9155 | 2023-09-12 | HE-301    | MED      | PM        | Bundle pulled for cleaning. Fouling factor    | Closed   | Patel S
                                                           above design. Cleaned with high-pressure water.
                                                           Performance restored. Schedule next clean
                                                           in 8 months.
WO#9312 | 2023-11-28 | V-302     | HIGH     | CM        | Pressure relief valve PSV-3021 found stuck    | Closed   | Mehta R
                                                           open during plant walk. Replaced PSV-3021.
                                                           Investigated: inlet accumulation of polymer
                                                           deposit. Recommend monthly functional test.
WO#9412 | 2024-01-17 | P-101     | HIGH     | CM        | Third seal failure. LI-2204 interlock still   | Closed   | Rajan K
                                                           not installed (pending instrumentation budget).
                                                           D-204 at 28% when failure occurred. URGENT:
                                                           install LI-2204 before next startup per
                                                           RCA#2024-07 recommendation.
WO#9501 | 2024-03-05 | P-101S    | MED      | PM        | Standby pump alternation and function test.   | Closed   | Singh A
                                                           P-101S started and ran for 4 hours. No issues.
WO#9620 | 2024-05-20 | H-301     | HIGH     | CM        | Tube skin thermocouple TE-3015 failed. Replaced| Closed  | Sharma D
                                                           TE-3015. No process impact during repair.
WO#9788 | 2024-08-14 | P-101     | HIGH     | CM        | LI-2204 interlock FINALLY installed. Function  | Closed  | Instrumentation team
                                                           tested: low-level trip at 35%, alarm at 40%.
                                                           P-101 seal failures expected to reduce
                                                           significantly. Close out RCA#2024-07.
WO#9901 | 2024-10-02 | C-401     | HIGH     | PM        | Annual column inspection. Trays 12-14         | Closed   | Inspection team
                                                           showing erosion; replaced tray 13. Corrosion
                                                           inhibitor injection rate increased per
                                                           recommendation of corrosion engineer.
WO#9998 | 2024-11-30 | P-101     | LOW      | PM        | Routine annual PM overhaul. No seal or        | Closed   | Rajan K
                                                           bearing replacement required this year —
                                                           first clean overhaul since LI-2204 installed.
                                                           Confirms cavitation was root cause of prior
                                                           failures.

SUMMARY STATISTICS (P-101, 2023-2024):
  Corrective maintenance events: 4
  Preventive maintenance events: 1
  Mean Time Between Failures (before LI-2204): 5.1 months
  Mean Time Between Failures (after LI-2204):  > 14 months (ongoing)
  Total downtime cost (estimated): INR 42 Lakhs
"""

(OUT / "maintenance_log_unit2.txt").write_text(WO_LOG, encoding="utf-8")
print("[OK] maintenance_log_unit2.txt")


# ═══════════════════════════════════════════════════════════════════════════════
# 4. Incident Investigation Report (PDF)
# ═══════════════════════════════════════════════════════════════════════════════
INCIDENT_PAGES = [
    """\
INCIDENT INVESTIGATION REPORT
Report No: INC-2024-03   |   Classification: High Potential Near-Miss
Date of Incident: 14-Feb-2024   |   Location: Unit 2, Pump P-101 Area
Reported by: Shift Supervisor Anil Sharma   |   Report Date: 18-Feb-2024

1. INCIDENT DESCRIPTION
At 14:35 on 14-Feb-2024, maintenance technician Rajan K entered the P-101 pump pit to
inspect the mechanical seal (WO#9788) while a hot-work permit (HWP#2024-022) was active
for welding on the adjacent cooling water line at EL+10.5 m, approximately 8 metres from
the pump pit opening.

The technician's personal gas detector alarmed for LEL at 12% (above the 10% LEL threshold
specified in SOP-SAF-014). The technician evacuated immediately. No injury occurred.

Gas source investigation found a small flange weep on the P-101 discharge line (not yet
depressurised at the time the entry permit was issued — procedural non-compliance).
""",
    """\
2. IMMEDIATE CAUSE
Confined space entry was initiated before the discharge line was fully depressurised and
positively isolated. The Maintenance Work Permit (MWP#2024-045) was signed by the day-shift
Area Authority who had not verified isolation before handover to the incoming shift.

3. CONTRIBUTING FACTORS
  a. No SIMOPS check performed: hot-work permit HWP#2024-022 was already active in the
     adjacent area when CSEP was issued. SIMOPS coordinator was not informed.
  b. Shift handover gap: the Area Authority who issued the MWP was 15 minutes into handover
     when the permit was countersigned; isolation verification was assumed, not confirmed.
  c. Permit system siloed: CSEP and HWP are managed in separate paper-based systems with
     no cross-check mechanism.

4. ROOT CAUSE
Absence of an integrated permit-to-work system that enforces SIMOPS checks and prevents
incompatible permits from being active simultaneously in overlapping geographic zones.
""",
    """\
5. CORRECTIVE ACTIONS (< 7 days)
  CA-01: Revoke HWP#2024-022 immediately; reinstate only after SIMOPS review. [DONE]
  CA-02: Conduct toolbox talk on SIMOPS requirements with all shift teams. [DONE 16-Feb]
  CA-03: Add mandatory SIMOPS coordinator sign-off field to all CSEP forms. [DONE 20-Feb]

6. PREVENTIVE ACTIONS (Long-term)
  PA-01: Implement digital PTW system with geofenced SIMOPS conflict detection.
         Owner: Plant Manager. Target: Q3 2024.
  PA-02: Integrate CMMS work orders with PTW system to flag concurrent high-risk activities.
         Owner: IT / Maintenance Dept. Target: Q4 2024.
  PA-03: Annual SIMOPS competency assessment for Area Authorities.
         Owner: Safety Dept. Target: Ongoing from Mar-2024.

7. LESSONS LEARNED
  • Data present, action absent: gas detector readings and permit logs both existed —
    what was missing was an intelligence layer that connected them before entry.
  • Simultaneous paper-based CSEP and HWP systems cannot prevent SIMOPS conflicts.
  • OISD-GDN-192 SIMOPS requirements are not being enforced at permit issuance stage.

8. REGULATORY REFERENCE
  OISD-GDN-192 Section 5.1: Simultaneous Operations must be identified and controlled.
  Factories Act 1948 Sec 36A: Safety officers must certify hazardous process safety systems.
  SOP-SAF-014 Section 5: Mandatory SIMOPS declaration before confined space entry.
""",
]

doc2 = fitz.open()
for body in INCIDENT_PAGES:
    page = doc2.new_page(width=595, height=842)
    page.insert_textbox(
        fitz.Rect(50, 50, 545, 800), body,
        fontsize=10, fontname="helv", align=0
    )
doc2.save(OUT / "incident_INC-2024-03.pdf")
doc2.close()
print("[OK] incident_INC-2024-03.pdf")


# ═══════════════════════════════════════════════════════════════════════════════
# 5. Regulatory Reference Extract (OISD key clauses – plain text)
# ═══════════════════════════════════════════════════════════════════════════════
REGULATORY_TEXT = """\
REGULATORY REFERENCE EXTRACTS – PlantBrain Knowledge Base
Compiled for: Unit 2 Coker Feed Section

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
OISD-GDN-192: MANAGEMENT OF CHANGE & SIMULTANEOUS OPERATIONS
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Section 4.3 – Permit-to-Work System
  A formal, documented Permit-to-Work (PTW) system shall be implemented for the following
  high-risk activities: hot work, confined space entry, energy isolation (LOTO), height work,
  radiography, and excavation. The PTW system shall be audited annually by a qualified safety
  officer. No work may commence without a valid, signed permit.

Section 5.1 – Simultaneous Operations (SIMOPS)
  Before issuing any permit, the Area Authority shall consult the SIMOPS register to identify
  any concurrently active permits in overlapping geographic zones. Where incompatible operations
  are identified (e.g., hot work adjacent to confined space entry, or pressure testing adjacent
  to mechanical work), the Area Authority shall either sequence the activities or obtain written
  approval from the Plant Manager with documented risk controls. A SIMOPS coordinator role
  shall be designated on each shift.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
OISD-STD-116: SAFE HANDLING OF HAZARDOUS CHEMICALS
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Section 7 – Confined Space Entry
  Atmospheric testing is mandatory before and continuously during confined space entry:
    Oxygen:               19.5 % minimum, 23.5 % maximum
    Flammable gas:        < 10 % LEL (Lower Explosive Limit)
    H₂S:                  < 10 ppm TWA
    CO:                   < 35 ppm TWA
  Testing instruments shall be calibrated within the preceding 6 months and bump-tested
  on the day of use. Results shall be recorded on the CSEP form.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
FACTORIES ACT 1948 – RELEVANT SECTIONS
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Section 36 & 36A – Precautions against dangerous fumes
  No person shall be required to enter any chamber, tank, vat, pit, pipe, flue, or similar
  confined space in which dangerous fumes are likely to be present unless it is provided with
  a manhole, or other means of egress, and unless a responsible person has certified, after
  adequate inspection, that the place is safe for entry.

Section 41 – Safety Audit for Hazardous Processes
  Every occupier carrying on a hazardous process shall conduct a safety audit of the process
  at least once in every two years by a qualified safety auditor.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
PESO – STATIC EQUIPMENT INSPECTION
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Pressure vessels operating above 1 kg/cm² shall be inspected under the Indian Boiler
Regulations (IBR) schedule. Records of all inspections, including thickness readings and
weld condition, shall be maintained on site and available for inspection by the Chief Inspector
of Factories. Vessels approaching corrosion allowance shall be de-rated or replaced.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
OISD-STD-144: HOT WORK IN OIL & GAS INDUSTRY
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Section 4 – Hot Work Permit Requirements
  Every hot-work permit shall specify:
    a. Gas test result at the work location (must be < 10% LEL)
    b. Standby fire extinguisher (minimum 9 kg DCP) positioned at work site
    c. Fire watcher assigned and named on the permit
    d. Hot work to cease if LEL reading rises above 10% at any time
  Hot-work permits are valid for one shift only (maximum 8 hours) and must be renewed.
"""

(OUT / "regulatory_extracts_OISD_FA.txt").write_text(REGULATORY_TEXT, encoding="utf-8")
print("[OK] regulatory_extracts_OISD_FA.txt")

print("\nAll sample documents created in ./sample_docs/")
